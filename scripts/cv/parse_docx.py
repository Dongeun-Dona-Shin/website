#!/usr/bin/env python3
"""
Parse a .docx CV and write structured JSON to src/generated/cv-data.json.

Usage:
    python3 scripts/cv/parse_docx.py --input public/cv.docx --output src/generated/cv-data.json
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx is not installed. Run: pip3 install python-docx", file=sys.stderr)
    sys.exit(1)


# ---------------------------------------------------------------------------
# Section heading detection
# ---------------------------------------------------------------------------

SECTION_PATTERNS = {
    "employment":      re.compile(r"employment|positions?|academic\s+appointments?|work\s+experience|research.{0,5}work\s+experience", re.I),
    "education":       re.compile(r"education|degrees?", re.I),
    "research_areas":  re.compile(r"research\s+(areas?|interests?|fields?)", re.I),
    "publications":    re.compile(r"^publications?$|refereed|peer.reviewed|journal\s+articles?", re.I),
    "work_in_progress": re.compile(r"work\s+in\s+progress|working\s+papers?|manuscripts?\s+under|under\s+review|forthcoming", re.I),
    "teaching":        re.compile(r"teaching|courses?\s+taught|pedagogy", re.I),
    "invited_talks":   re.compile(r"invited\s+talks?|invited\s+presentations?|^presentations?$|^conferences?$", re.I),
    "honors_grants":   re.compile(r"honors?|^grants?$|^awards?$|fellowships?|funding|^scholarship$", re.I),
    "service":         re.compile(r"service|professional\s+activities?|reviewing|committees?", re.I),
}


def classify_heading(text: str) -> str | None:
    t = text.strip()
    if not t:
        return None
    for key, pat in SECTION_PATTERNS.items():
        if pat.search(t):
            return key
    return None


def is_heading_style(para) -> bool:
    style_name = (para.style.name or "").lower()
    return "heading" in style_name or style_name.startswith("title")


def para_is_bold_line(para) -> bool:
    """True if every non-empty run in the paragraph is bold."""
    runs = [r for r in para.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)


def para_is_allcaps_heading(text: str) -> bool:
    """True if the line looks like an ALL-CAPS section header (e.g. PUBLICATIONS, SCHOLARSHIP)."""
    t = text.strip()
    if len(t) < 3 or len(t) > 80:
        return False
    # Must be mostly uppercase letters (allow spaces, punctuation, middle-dot ·∙•)
    letters = [c for c in t if c.isalpha()]
    return bool(letters) and sum(1 for c in letters if c.isupper()) / len(letters) >= 0.85


# ---------------------------------------------------------------------------
# Text extraction helpers
# ---------------------------------------------------------------------------

def full_text(para) -> str:
    return "".join(r.text for r in para.runs).strip()


def extract_name_and_contact(doc) -> dict:
    """Best-effort: first 1–3 non-empty paragraphs often contain name/contact."""
    result = {"name": "", "email": "", "phone": "", "website": "", "affiliation": ""}
    lines = []
    for para in doc.paragraphs:
        t = full_text(para)
        if t:
            lines.append(t)
        if len(lines) >= 6:
            break
    if lines:
        result["name"] = lines[0]
    for line in lines[1:]:
        if "@" in line and not result["email"]:
            result["email"] = line
        elif re.search(r"http|www\.", line, re.I) and not result["website"]:
            result["website"] = line
        elif re.search(r"\d{3}[-.\s]\d{3,4}", line) and not result["phone"]:
            result["phone"] = line
        elif not result["affiliation"] and line != result["name"]:
            result["affiliation"] = line
    return result


# ---------------------------------------------------------------------------
# Main parser
# ---------------------------------------------------------------------------

def parse_cv(path: Path) -> dict:
    doc = Document(str(path))

    data: dict = {
        "meta": extract_name_and_contact(doc),
        "employment": [],
        "education": [],
        "research_areas": [],
        "publications": [],
        "work_in_progress": [],
        "teaching": [],
        "invited_talks": [],
        "honors_grants": [],
        "service": [],
        "_unparsed": [],
    }

    current_section: str | None = None
    buffer: list[str] = []

    def flush(section: str, lines: list[str]):
        if not lines:
            return
        text = "\n".join(lines).strip()
        if not text:
            return
        if section in data and isinstance(data[section], list):
            data[section].append(text)
        else:
            data["_unparsed"].append({"section": section, "text": text})

    for para in doc.paragraphs:
        text = full_text(para)
        if not text:
            continue

        is_heading = is_heading_style(para) or para_is_bold_line(para) or para_is_allcaps_heading(text)

        if is_heading:
            key = classify_heading(text)
            if key:
                flush(current_section, buffer)
                buffer = []
                current_section = key
                continue
            # Bold line that doesn't match a known section — treat as content
        # If we haven't found any section yet, skip header lines
        if current_section is None:
            continue

        buffer.append(text)

    flush(current_section, buffer)

    # Post-process: split multi-line buffers into individual entries
    for key in ["employment", "education", "publications", "work_in_progress",
                "teaching", "invited_talks", "honors_grants", "service"]:
        merged: list[str] = []
        for block in data[key]:
            # Each paragraph already came in as a separate entry; keep them as-is
            merged.append(block)
        data[key] = merged

    # research_areas: split on semicolons or newlines into a flat list
    flat_areas: list[str] = []
    for block in data["research_areas"]:
        for item in re.split(r";|\n|•|·", block):
            item = item.strip()
            if item:
                flat_areas.append(item)
    data["research_areas"] = flat_areas

    return data


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Parse a .docx CV into JSON.")
    parser.add_argument("--input", default="public/cv.docx", help="Path to the .docx CV file")
    parser.add_argument("--output", default="src/generated/cv-data.json", help="Output JSON path")
    args = parser.parse_args()

    input_path = Path(args.input)
    # Fallback: if explicit path not found, try root cv.docx for backwards compat
    if not input_path.exists() and args.input == "public/cv.docx":
        fallback = Path("cv.docx")
        if fallback.exists():
            print(f"INFO: public/cv.docx not found; using root cv.docx as fallback", file=sys.stderr)
            input_path = fallback

    if not input_path.exists():
        print(f"ERROR: CV file not found at '{input_path}'. "
              f"Place your CV at public/cv.docx and re-run.", file=sys.stderr)
        sys.exit(1)

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    print(f"Parsing: {input_path}")
    data = parse_cv(input_path)

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    print(f"Written: {output_path}")
    print(f"  Sections: {', '.join(k for k in data if k not in ('meta', '_unparsed') and data[k])}")
    if data["_unparsed"]:
        print(f"  Unparsed blocks: {len(data['_unparsed'])} (check _unparsed in JSON)")


if __name__ == "__main__":
    main()
